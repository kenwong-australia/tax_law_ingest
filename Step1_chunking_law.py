#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Legislation Document Processor with TOC Extraction and Fuzzy Header-Based Chunking

- Loads DOCX files from a specified directory.
- Extracts and cleans the TOC from each DOCX.
- Uses RapidFuzz to fuzzy-match headers in the main document.
- Chunks the document by header and writes chunk details to individual JSON files.
- Saves a summary log file for each DOCX processed showing:
    * Total TOC entries,
    * Headers matched (out of the TOC total),
    * Number of chunks (JSON files) produced,
    * Overall processing time.
"""

import re
import os
import time
from pathlib import Path
import docx
from datetime import datetime
from rapidfuzz import fuzz
import json
import openai
import logging
import sys
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Set OpenAI API key from environment variable
openai.api_key = os.environ.get("OPENAI_API_KEY")
if not openai.api_key:
    raise ValueError("OPENAI_API_KEY environment variable not found. Please check your .env file.")

# Load keyword library
with open('tax_keywords.json', 'r', encoding='utf-8') as kw_file:
    keyword_dict = json.load(kw_file)

# Create mappings for case-insensitive matching
KEYWORD_TO_CATEGORY = {}
KEYWORD_ORIGINAL_CASE = {}

for category, keywords in keyword_dict.items():
    category_lower = category.lower()
    KEYWORD_ORIGINAL_CASE[category_lower] = category
    KEYWORD_TO_CATEGORY[category_lower] = category
    for keyword in keywords:
        keyword_lower = keyword.lower()
        KEYWORD_ORIGINAL_CASE[keyword_lower] = keyword
        KEYWORD_TO_CATEGORY[keyword_lower] = category

# Define TAX_KEYWORDS as a flat list of all keywords and categories from the keyword library
TAX_KEYWORDS = list(KEYWORD_ORIGINAL_CASE.values())

# --- CONFIGURATION ---
# Create local log directory
LOG_DIR = "/Users/kenmacpro/pinecone-upsert/logs/chunking"
os.makedirs(LOG_DIR, exist_ok=True)

# Set up logging with timestamp
timestamp_str = datetime.now().strftime("%Y%m%d_%H%M%S")
MAIN_LOG_FILE_NAME = os.path.join(LOG_DIR, f"log_chunking_law_{timestamp_str}.txt")

# Log the current working directory
print(f"Current working directory: {os.getcwd()}")
print(f"Log files will be saved to: {LOG_DIR}")

# Configure the root logger only once at the beginning
logger = logging.getLogger()
logger.setLevel(logging.INFO)
formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')

# Console handler - add only once
console_handler = logging.StreamHandler(sys.stdout)
console_handler.setFormatter(formatter)
logger.addHandler(console_handler)

# Main file handler
main_file_handler = logging.FileHandler(MAIN_LOG_FILE_NAME)
main_file_handler.setFormatter(formatter)
logger.addHandler(main_file_handler)

# Global debug flag.
DEBUG = True

def debug_print(message):
    if DEBUG:
        print(message)

LOCAL_DIR = "/Users/kenmacpro/pinecone-upsert/testfiles_law"

def get_all_docx_files(directory):
    """Return all DOCX files found in the directory."""
    p = Path(directory)
    docx_files = list(p.glob("*.docx"))
    return docx_files

def extract_toc(doc):
    """
    Extract TOC paragraphs from the DOCX using paragraph styles.
    Focuses on 'toc 5' style entries.
    Returns a tuple of (toc_paragraphs, end_index) where end_index is the paragraph index after the TOC.
    """
    toc_paragraphs = []
    toc_start_index = -1
    toc_end_index = 0
    
    # Look for paragraphs with style 'toc 5'
    for i, p in enumerate(doc.paragraphs):
        style_name = p.style.name.lower()
        if style_name == 'toc 5':  # Filter specifically for TOC 5
            if toc_start_index == -1:
                toc_start_index = i
            toc_paragraphs.append(p.text)
            toc_end_index = i + 1  # Next paragraph index
    
    # Fallback to any TOC style if none found
    if not toc_paragraphs:
        print("No TOC 5 styles found. Trying any TOC styles as fallback.")
        for i, p in enumerate(doc.paragraphs):
            style_name = p.style.name.lower()
            if 'toc' in style_name:
                if toc_start_index == -1:
                    toc_start_index = i
                toc_paragraphs.append(p.text)
                toc_end_index = i + 1
    
    # Fallback to text-based TOC detection
    if not toc_paragraphs:
        print("No TOC styles found. Falling back to text-based TOC detection.")
        found_contents = False
        for i, p in enumerate(doc.paragraphs):
            text = p.text.strip()
            if not found_contents:
                if "contents" in text.lower():
                    found_contents = True
                    toc_paragraphs.append(text)
                    toc_start_index = i
                    toc_end_index = i + 1
            else:
                if p.style.name.lower().startswith("heading"):
                    toc_end_index = i
                    break
                if text:
                    toc_paragraphs.append(text)
                    toc_end_index = i + 1

    print(f"Found TOC with {len(toc_paragraphs)} paragraphs, starting at {toc_start_index}, ending at {toc_end_index}")
    
    # Print a sample for verification
    if toc_paragraphs:
        print("Sample TOC entries:")
        for entry in toc_paragraphs[:5]:
            print(f"  {entry}")
    
    return toc_paragraphs, toc_end_index

def clean_and_parse_toc_from_doc(doc):
    """Extract and clean TOC entries from a document."""
    toc_entries = []
    in_toc = False
    toc_text = []
    
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text and (text.lower().startswith("table of contents") or text.lower() == "contents"):
            in_toc = True
            continue
        
        if in_toc:
            if not text:  # Skip empty lines
                continue
                
            # Check for patterns that indicate end of TOC
            if (text.startswith("Chapter") and len(text) < 30) or text == "Part 1":
                in_toc = False
                break
                
            # If line looks like a TOC entry, add it
            toc_text.append(text)
    
    # Updated pattern to be more flexible with section number formats
    # Simplified to not extract page numbers and allow more variations in section format
    section_pattern = r'^\s*(?:(\d+(?:[.-]\d+)*[A-Z]*)|\s*(\d+[A-Z]+))\s+(.*?)(?:\s*\.{2,}|\s*\t|\s+\d+)?$'
    skip_keywords = ("part", "division", "subdivision", "chapter", "guide", "operative provisions")
    
    # Process each line in the TOC
    section_numbers_seen = set()  # Track section numbers to avoid duplicates
    
    for line in toc_text:
        # Skip organizational elements like "Part 1" etc.
        if any(line.lower().startswith(keyword) for keyword in skip_keywords):
            continue
        
        # Filter out definition entries (they typically have a long explanation after the number)
        if len(line.split("\t")) > 2 and len(line) > 100:  # Long line with tabs is likely a definition
            continue
        
        # Also filter out lines that start with numbers but aren't proper section entries
        # (They typically contain words like "means" that indicate definitions)
        if "means" in line or "in relation to" in line:
            continue
        
        # Skip sections with the title 'What this Division is about'
        if 'what this division is about' in line.lower():
            continue
        
        match = re.match(section_pattern, line)
        if match:
            section1, section2, title = match.groups()[:3]  # Only get first 3 groups, ignore page number
            section_number = section1 if section1 else section2
            
            # Skip duplicate section numbers
            if section_number in section_numbers_seen:
                continue
                
            section_numbers_seen.add(section_number)
            
            # Clean the title of tabs, extra whitespace, and page numbers
            clean_title = re.sub(r'\s+', ' ', title.strip())
            clean_title = re.sub(r'\s*\.{2,}.*$', '', clean_title)  # Remove dots and page numbers
            
            toc_entries.append({
                'section_number': section_number,
                'title': clean_title,
                'page_number': None  # No longer extracting page numbers
            })
    
    return toc_entries

def write_toc_entries(toc_entries, output_file):
    """Write cleaned TOC entries to a file for inspection."""
    with open(output_file, "w", encoding="utf-8") as f:
        for entry in toc_entries:
            # Include page number in the output file for debugging
            page_info = f" (Page {entry['page_number']})" if entry['page_number'] else ""
            f.write(f'{entry["section_number"]}\t{entry["title"]}{page_info}\n')

def find_header_paragraph_index(doc, start_index, header, threshold=60):
    """Find the paragraph index in the document that matches the given header."""
    debug_print(f"Searching for header: '{header}'")
    
    # More flexible header splitting - handle various spacing
    parts = [p.strip() for p in re.split(r'\s{2,}', header, maxsplit=1)]
    if len(parts) != 2:
        debug_print(f"Header format is unexpected: '{header}'")
        return -1
        
    section_number, title = parts[0].strip(), parts[1].strip()
    debug_print(f"Looking for section '{section_number}' with title '{title}'")
    
    # Clean section number for comparison
    clean_section = re.sub(r'\s+', '', section_number)
    
    for i in range(start_index, len(doc.paragraphs)):
        p = doc.paragraphs[i]
        text = p.text.strip()
        
        # Clean paragraph text section number for comparison
        text_clean = re.sub(r'\s+', '', text)
        
        # More flexible section number matching
        if (text_clean.startswith(clean_section) or 
            clean_section in text_clean[:len(clean_section) + 2]):  # Allow slight variation
            debug_print(f"Found section number match at paragraph {i}: '{text}'")
            header_parts = re.split(r'\s{2,}', text)
            if len(header_parts) >= 2:
                found_title = header_parts[1]
                title_score = fuzz.ratio(title.lower(), found_title.lower())
                debug_print(f"Title match score: {title_score} ('{title}' vs '{found_title}')")
                if title_score >= threshold:
                    debug_print(f"✓ MATCH: Section {section_number} at paragraph {i}")
                    return i
            # Only use section number as fallback if we're really confident about the match
            elif text_clean == clean_section:
                debug_print(f"Found exact section number match at paragraph {i}")
                return i
    
    debug_print(f"✗ No match found for section '{section_number}'")
    return -1

def text_to_markdown(paragraphs, start_idx, end_idx, section_number, section_title):
    """Convert document paragraphs to Markdown format."""
    markdown_lines = []
    first_para = paragraphs[start_idx].text.strip() if start_idx < len(paragraphs) else ""
    if first_para.startswith(section_number):
        markdown_lines.append(f"# {first_para}\n")
        start_process_idx = start_idx + 1
    else:
        markdown_lines.append(f"# {section_number} {section_title}\n")
        start_process_idx = start_idx
    
    current_para = []
    for i in range(start_process_idx, end_idx):
        p = paragraphs[i]
        text = p.text.strip()
        if text:
            is_subsection = (
                re.match(r'^\(\d+\)', text) or
                re.match(r'^\d+\.\d+', text) or
                re.match(r'^[A-Z]\.\s+', text)
            )
            if is_subsection:
                if current_para:
                    markdown_lines.append(" ".join(current_para) + "\n\n")
                    current_para = []
                markdown_lines.append(f"## {text}\n\n")
            else:
                current_para.append(text)
        else:
            if current_para:
                markdown_lines.append(" ".join(current_para) + "\n\n")
                current_para = []
    
    if current_para:
        markdown_lines.append(" ".join(current_para) + "\n")
    
    return "".join(markdown_lines)

def extract_keywords_with_llm(text, keyword_library, section_number=None, source_file=None):
    # Log section and source information
    if section_number and source_file:
        logging.info(f"\nProcessing Section {section_number} from {source_file}")
    
    prompt = f"""
    Extract relevant keywords from the following text based on the provided keyword library.

    Text:
    {text}

    Keyword Library:
    {', '.join(keyword_library)}

    Return keywords as a comma-separated list.
    """
    
    max_retries = 3
    retry_delay = 5  # seconds
    
    for attempt in range(max_retries):
        try:
            # Attempt to call the OpenAI API
            response = openai.ChatCompletion.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.1,
                max_tokens=100,
                request_timeout=60  # Set a timeout of 60 seconds
            )
            
            # Extract keywords from the response with better error handling
            try:
                # Log the response type for debugging
                logging.debug(f"Response type: {type(response)}")
                
                # Access the content based on the response type
                try:
                    if hasattr(response, 'choices') and len(response.choices) > 0:
                        # Using attribute access for OpenAIObject
                        if hasattr(response.choices[0], 'message') and hasattr(response.choices[0].message, 'content'):
                            keywords_text = response.choices[0].message.content
                            if keywords_text is not None:
                                keywords_text = keywords_text.strip()
                            else:
                                keywords_text = ""
                        else:
                            raise AttributeError("Message or content attribute not found in response")
                    elif isinstance(response, dict) and 'choices' in response and len(response['choices']) > 0:
                        # Dictionary access
                        keywords_text = response['choices'][0]['message']['content'].strip()
                    else:
                        raise ValueError("Unexpected response structure")
                        
                except (AttributeError, IndexError, KeyError) as access_error:
                    # Handle common access errors
                    logging.warning(f"Error accessing response content: {str(access_error)}")
                    # Try converting to string and parsing
                    try:
                        response_str = str(response)
                        if '"content": "' in response_str:
                            content_start = response_str.find('"content": "') + 12
                            content_end = response_str.find('"', content_start)
                            keywords_text = response_str[content_start:content_end].strip()
                        else:
                            logging.error("Could not find content in response string")
                            return [], []
                    except Exception as e:
                        logging.error(f"Failed to extract content from response string: {str(e)}")
                        return [], []
                
                # Check if we actually got any keywords
                if not keywords_text:
                    logging.warning("No keywords text found in the response")
                    return [], []
                    
                # Log the raw keywords response for debugging
                logging.info(f"Raw keywords response: {keywords_text}")
                
                # Parse the comma-separated keywords
                raw_keywords = [k.strip() for k in keywords_text.split(',') if k.strip()]
                logging.info(f"Raw keywords after splitting: {raw_keywords}")
                
                # Case-insensitive keyword matching
                # Convert all library keywords to lowercase for comparison
                keyword_library_lower = [k.lower() for k in keyword_library]
                valid_keywords = set()
                matched_categories = set()
                
                for keyword in raw_keywords:
                    keyword_lower = keyword.lower()

                    if keyword_lower in KEYWORD_TO_CATEGORY:
                        original_keyword = KEYWORD_ORIGINAL_CASE[keyword_lower]
                        category = KEYWORD_TO_CATEGORY[keyword_lower]
                        valid_keywords.add(original_keyword)
                        matched_categories.add(category)
                        logging.info(f"Matched keyword '{original_keyword}' under category '{category}'")
                    else:
                        # Partial matching fallback
                        for lib_keyword_lower in KEYWORD_TO_CATEGORY:
                            if keyword_lower in lib_keyword_lower or lib_keyword_lower in keyword_lower:
                                original_keyword = KEYWORD_ORIGINAL_CASE[lib_keyword_lower]
                                category = KEYWORD_TO_CATEGORY[lib_keyword_lower]
                                valid_keywords.add(original_keyword)
                                matched_categories.add(category)
                                logging.info(f"Partial match '{keyword}' → '{original_keyword}' under category '{category}'")
                                break
                
                # Convert sets to sorted lists for consistent ordering
                valid_keywords = sorted(valid_keywords)
                matched_categories = sorted(matched_categories)
                
                logging.info(f"FINAL KEYWORDS: {len(valid_keywords)} keywords: {valid_keywords}")
                return valid_keywords, matched_categories
            except Exception as parsing_error:
                # More detailed error logging for response parsing issues
                logging.error(f"Error parsing OpenAI response: {str(parsing_error)}")
                logging.error(f"Response type: {type(response)}")
                logging.error(f"Response content: {str(response)[:500]}...")  # Log first 500 chars
                return [], []
            
        except (openai.error.Timeout, openai.error.APIError, openai.error.ServiceUnavailableError) as e:
            # Log the error
            if attempt < max_retries - 1:
                logging.warning(f"OpenAI API error (attempt {attempt+1}/{max_retries}): {str(e)}. Retrying in {retry_delay} seconds...")
                time.sleep(retry_delay)
                # Increase the delay for next retry (exponential backoff)
                retry_delay *= 2
            else:
                logging.error(f"Failed to extract keywords after {max_retries} attempts: {str(e)}")
                # Return an empty list of keywords if all retries fail
                return [], []
        
        except Exception as e:
            # Handle any other unexpected errors
            logging.error(f"Unexpected error in keyword extraction: {str(e)}")
            return [], []

def chunk_document_by_toc_paragraphs(doc, toc_entries, main_start_index, legislation_title,
                                     docx_filename=None, current_datetime=None):
    """
    Chunk the document by scanning paragraphs for fuzzy-matched headers,
    prepending the user-provided legislation_title to each chunk.
    """
    debug_print(f"Chunking document: {len(toc_entries)} TOC entries starting at paragraph {main_start_index}")
    
    if main_start_index > len(doc.paragraphs) / 2:
        print(f"WARNING: TOC end index {main_start_index} seems too high. Using index 100 as fallback.")
        main_start_index = 100
    
    header_indices = []
    for entry in toc_entries:
        expected_header = f"{entry['section_number']}  {entry['title']}"
        debug_print(f"Searching for header: '{expected_header}'")
        idx = find_header_paragraph_index(doc, main_start_index, expected_header, threshold=60)
        debug_print(f"Header search result: {idx}")
        if idx != -1:
            header_indices.append((idx, entry))
    
    # Add this after finding all header indices but before sorting them
    # Remove duplicate section headers that are too close to each other
    filtered_header_indices = []
    section_positions = {}

    for idx, entry in header_indices:
        section_number = entry['section_number']
        
        if section_number in section_positions:
            # If we already found this section, check how close they are
            prev_idx = section_positions[section_number]
            distance = idx - prev_idx
            
            if distance < 10:  # If headers are less than 10 paragraphs apart
                print(f"Warning: Ignoring duplicate header for section {section_number} at position {idx} (too close to previous at {prev_idx})")
                continue
            else:
                # Keep both but log the duplicate
                print(f"Note: Found another instance of section {section_number} at position {idx}, keeping both (distance: {distance} paragraphs)")
                
        # Store this position and keep this header
        section_positions[section_number] = idx
        filtered_header_indices.append((idx, entry))

    # Replace the original list with filtered one
    header_indices = filtered_header_indices
    
    header_indices.sort(key=lambda x: x[0])
    chunks = []
    for i, (start_idx, entry) in enumerate(header_indices):
        if i == len(header_indices) - 1:
            for j in range(start_idx + 1, len(doc.paragraphs)):
                text = doc.paragraphs[j].text.strip()
                if re.match(r'^\s*(?:\d+-\d+|\d+[A-Z]+)\s+\w+', text):
                    end_idx = j
                    break
                if j > start_idx + 500:
                    end_idx = start_idx + 500
                    break
            else:
                end_idx = len(doc.paragraphs)
        else:
            end_idx = header_indices[i + 1][0]
        
        markdown_text = text_to_markdown(doc.paragraphs, start_idx, end_idx,
                                         entry['section_number'], entry['title'])
        
        section_with_prefix = f"{legislation_title} {entry['section_number']} {entry['title']}"
        full_reference = f"{legislation_title} {entry['section_number']} {entry['title']}"
        
        # Measure LLM processing time
        start_llm_time = time.time()

        # Extract keywords using LLM
        keywords, categories = extract_keywords_with_llm(markdown_text, TAX_KEYWORDS, entry['section_number'], docx_filename)

        # Calculate and print LLM processing time
        llm_processing_time = time.time() - start_llm_time
        print(f"LLM processing time for keywords extraction: {llm_processing_time:.2f} seconds")

        # Adjust section to legislation + section numbering only
        section_number_only = entry['section_number']

        # Prepend full reference to text
        markdown_text_with_ref = f"{full_reference}\n\n{markdown_text}"

        # Add section number to keywords list
        keywords_with_section = list(keywords)
        if section_number_only not in keywords_with_section:
            keywords_with_section.append(section_number_only)
            logging.info(f"Added section number '{section_number_only}' to keywords")

        chunk_metadata = {
            "keywords": keywords_with_section,
            "categories": list(categories),
            "full_reference": full_reference,
            "source_file": docx_filename or "unknown",
            "creation_date": current_datetime or datetime.now().isoformat()
        }

        chunks.append({
            "section": section_number_only,
            "text": markdown_text_with_ref,
            "metadata": chunk_metadata
        })
    
    # Add debug but remove CSV generation code
    print(f"Debug: chunks list has {len(chunks)} items")
    for i, chunk in enumerate(chunks[:3]):  # Print first 3 for verification
        print(f"Debug: Chunk {i+1} section: {chunk['section']}")
    
    # Return the chunks without writing CSV
    return chunks

def write_individual_summary_log(log_file, docx_filename, toc_total, matches, chunks_count, time_taken, chunks=None):
    """Write an individual summary log for a DOCX file."""
    # Convert time_taken to hours, minutes, and seconds
    hours, rem = divmod(time_taken, 3600)
    minutes, seconds = divmod(rem, 60)
    
    # Calculate unmatched and percentage
    unmatched = toc_total - matches
    match_percentage = (matches / toc_total * 100) if toc_total > 0 else 0
    
    with open(log_file, 'w', encoding='utf-8') as f:
        f.write(f"Processing Summary for {docx_filename}\n")
        f.write("=" * 40 + "\n")
        f.write(f"Total TOC entries: {toc_total}\n")
        f.write(f"Headers Matched: {matches} out of {toc_total} ({match_percentage:.1f}%)\n")
        f.write(f"Headers NOT Matched: {unmatched}\n")
        f.write(f"Chunks produced: {chunks_count}\n")
        f.write(f"Time taken: {int(hours)}h {int(minutes)}m {int(seconds)}s\n")
        
        # Only calculate keywords if chunks are provided
        if chunks:
            f.write(f"Keywords extracted: {sum(len(chunk['metadata']['keywords']) for chunk in chunks)}\n")
        else:
            f.write("Keywords extracted: Not calculated\n")

def write_chunks_to_csv(chunks, output_file):
    """Write all chunks to a CSV file."""
    import csv
    chunk_count = len(chunks)
    print(f"Writing {chunk_count} chunks to CSV file: {output_file}")
    
    if chunk_count == 0:
        print("Warning: No chunks to write to CSV!")
        return False
        
    try:
        rows_written = 0
        with open(output_file, 'w', newline='', encoding='utf-8') as csvfile:
            fieldnames = ['section', 'text', 'full_reference', 'source_file', 'creation_date', 'keywords']
            writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
            writer.writeheader()
            
            for chunk in chunks:
                try:
                    row = {
                        'section': chunk['section'],
                        'text': chunk['text'],
                        'full_reference': chunk['metadata']['full_reference'],
                        'source_file': chunk['metadata']['source_file'],
                        'creation_date': chunk['metadata']['creation_date'],
                        'keywords': ', '.join(chunk['metadata']['keywords'])
                    }
                    writer.writerow(row)
                    rows_written += 1
                except Exception as chunk_error:
                    print(f"Error writing chunk to CSV: {chunk_error}")
                    
        print(f"Successfully wrote {rows_written} out of {chunk_count} chunks to CSV file")
        return rows_written > 0
    except Exception as e:
        print(f"Error writing CSV file: {str(e)}")
        return False

def setup_file_logger(docx_filename, overall_dt_string):
    """Set up a specific logger for a DOCX file."""
    # Remove any existing file handlers (except main handler)
    logger = logging.getLogger()
    for handler in logger.handlers[:]:
        if isinstance(handler, logging.FileHandler) and handler.baseFilename != MAIN_LOG_FILE_NAME:
            logger.removeHandler(handler)
    
    # Create a new file handler for this DOCX file
    log_file_name = os.path.join(LOG_DIR, f"log_{docx_filename}_{overall_dt_string}.txt")
    file_handler = logging.FileHandler(log_file_name)
    file_handler.setFormatter(logging.Formatter('%(asctime)s - %(levelname)s - %(message)s'))
    logger.addHandler(file_handler)
    
    return log_file_name

def main():
    # Track overall start time
    overall_start_time = time.time()
    
    # Print welcome banner
    print("\n" + "=" * 80)
    print(" " * 30 + "LAW DOCUMENT CHUNKING PIPELINE")
    print("=" * 80 + "\n")
    
    logging.info("Starting the law document chunking pipeline.")
    
    # Directory containing the JSON files and checkpoint file
    checkpoint_path = os.path.join(LOG_DIR, "chunking_law_checkpoint.txt")
    processed_files = []  # Track successfully processed files
    failed_files = []     # Track failed files
    time_per_file = {}    # Track processing time for each file
    
    # Ask user if they want to reset the checkpoint
    reset_checkpoint = input("Do you want to reset the checkpoint? (yes/no): ").strip().lower()
    if reset_checkpoint == 'yes':
        start_index = 0
        logging.info("Checkpoint reset to 0.")
    else:
        # Read checkpoint if it exists
        start_index = 0
        if os.path.exists(checkpoint_path):
            try:
                with open(checkpoint_path, "r") as cp:
                    start_index = int(cp.read().strip())
                    logging.info(f"Resuming from checkpoint: file index {start_index}")
            except Exception as e:
                logging.error(f"Error reading checkpoint file: {e}")

    # Ask if user wants to run in test mode (process only first 20 chunks)
    test_mode = input("Run in test mode (process only first 20 chunks)? (yes/no): ").strip().lower() == 'yes'
    if test_mode:
        logging.info("TEST MODE ACTIVE: Processing only the first 20 chunks.")
        print("\n*** TEST MODE ENABLED - Only processing first 20 chunks ***\n")

    try:
        docx_files = get_all_docx_files(LOCAL_DIR)
        if not docx_files:
            print("No DOCX files found in the directory.")
            return
        
        print(f"Found {len(docx_files)} DOCX files to process")
        json_dir = os.path.join(LOCAL_DIR, "json")
        os.makedirs(json_dir, exist_ok=True)
        overall_dt_string = datetime.now().strftime('%Y%m%d_%H%M%S')
        
        for docx_index, docx_path in enumerate(docx_files, 1):
            file_start = time.time()  # Start timing for this DOCX file
            
            # Calculate elapsed time since start of run
            elapsed_since_start = time.time() - overall_start_time
            hours, rem = divmod(elapsed_since_start, 3600)
            minutes, seconds = divmod(rem, 60)
            
            print(f"\n==== Processing DOCX file: {docx_path.name} ({docx_index}/{len(docx_files)}) ====")
            print(f"Time elapsed since start: {int(hours)}h {int(minutes)}m {int(seconds)}s")
            docx_filename = docx_path.name
            
            # Set up a unique log file for each DOCX file
            log_file_name = setup_file_logger(docx_filename, overall_dt_string)
            logging.info(f"Processing file: {docx_filename}")
            
            try:
                # Open the document
                doc = docx.Document(str(docx_path))
                
                # Extract legislation title from filename or first line
                legislation_title = os.path.basename(docx_path).split('.')[0]
                # Try to get a better title from the first few paragraphs
                for i in range(min(5, len(doc.paragraphs))):
                    if doc.paragraphs[i].text and len(doc.paragraphs[i].text.strip()) > 10:
                        legislation_title = doc.paragraphs[i].text.strip()
                        break
                
                logging.info(f"Legislation title: {legislation_title}")
                
                toc_entries = clean_and_parse_toc_from_doc(doc)
                toc_total = len(toc_entries)
                print(f"Extracted {toc_total} TOC entries")
                
                toc_output_file = os.path.join(LOCAL_DIR, f"TOC_clean_{docx_path.stem}.txt")
                write_toc_entries(toc_entries, toc_output_file)
                print(f"Cleaned TOC written to {toc_output_file}")
                
                toc_paragraphs, main_start_index = extract_toc(doc)
                print(f"TOC ends at paragraph {main_start_index}")
                for i in range(main_start_index, min(main_start_index + 10, len(doc.paragraphs))):
                    print(f"Paragraph {i}: '{doc.paragraphs[i].text}'")
                
                # Fuzzy match TOC entries in batches
                match_counter = 0
                header_indices = []
                current_search_start = main_start_index
                batch_size = 50
                for batch_start in range(0, toc_total, batch_size):
                    batch_end = min(batch_start + batch_size, toc_total)
                    
                    # Calculate elapsed time for batch update
                    batch_elapsed = time.time() - overall_start_time
                    b_hours, b_rem = divmod(batch_elapsed, 3600)
                    b_minutes, b_seconds = divmod(b_rem, 60)
                    
                    print(f"\nProcessing batch {batch_start//batch_size + 1}: entries {batch_start+1}-{batch_end}")
                    print(f"Time elapsed since start: {int(b_hours)}h {int(b_minutes)}m {int(b_seconds)}s")
                    for idx, entry in enumerate(toc_entries[batch_start:batch_end], start=batch_start):
                        # Calculate elapsed time since start of run for this entry
                        entry_elapsed = time.time() - overall_start_time
                        e_hours, e_rem = divmod(entry_elapsed, 3600)
                        e_minutes, e_seconds = divmod(e_rem, 60)
                        
                        expected_header = f"{entry['section_number']}  {entry['title']}"
                        print(f"Matching TOC entry {idx + 1}/{toc_total}: '{expected_header}' (starting at para {current_search_start}) [Time: {int(e_hours)}h {int(e_minutes)}m {int(e_seconds)}s]", end="", flush=True)
                        idx_found = find_header_paragraph_index(doc, current_search_start, expected_header, threshold=60)
                        if idx_found != -1:
                            match_counter += 1
                            print(f" ✓ MATCH at paragraph {idx_found}")
                            header_indices.append((idx_found, entry))
                            current_search_start = idx_found + 1
                        else:
                            # Calculate time again for retry attempt
                            retry_elapsed = time.time() - overall_start_time
                            r_hours, r_rem = divmod(retry_elapsed, 3600)
                            r_minutes, r_seconds = divmod(r_rem, 60)
                            
                            print(f" ✗ NOT FOUND; retrying from TOC end... [Time: {int(r_hours)}h {int(r_minutes)}m {int(r_seconds)}s]", end="", flush=True)
                            idx_found = find_header_paragraph_index(doc, main_start_index, expected_header, threshold=60)
                            if idx_found != -1:
                                match_counter += 1
                                print(f" ✓ MATCH at paragraph {idx_found}")
                                header_indices.append((idx_found, entry))
                            else:
                                print(" ✗ STILL NOT FOUND")
                    print(f"Batch complete: Matched {match_counter} out of {batch_end} entries so far")
                    print(f"Current search position: paragraph {current_search_start}")
                
                print(f"\nFINAL SUMMARY: Matched {match_counter} out of {toc_total} total sections")
                
                # Log summary of TOC matching to the per-DOCX log file
                logging.info(f"TOC MATCHING SUMMARY: Successfully matched {match_counter} out of {toc_total} TOC entries ({(match_counter/toc_total)*100:.1f}%)")
                
                header_indices.sort(key=lambda x: x[0])
                chunks = []
                for i, (start_idx, entry) in enumerate(header_indices):
                    if i == len(header_indices) - 1:
                        for j in range(start_idx + 1, len(doc.paragraphs)):
                            text = doc.paragraphs[j].text.strip()
                            if re.match(r'^\s*(?:\d+-\d+|\d+[A-Z]+)\s+\w+', text):
                                end_idx = j
                                break
                            if j > start_idx + 500:
                                end_idx = start_idx + 500
                                break
                        else:
                            end_idx = len(doc.paragraphs)
                    else:
                        end_idx = header_indices[i + 1][0]
                    
                    markdown_text = text_to_markdown(doc.paragraphs, start_idx, end_idx,
                                                     entry['section_number'], entry['title'])
                    section_with_prefix = f"{legislation_title} {entry['section_number']} {entry['title']}"
                    full_reference = f"{legislation_title} {entry['section_number']} {entry['title']}"
                    
                    # Measure LLM processing time
                    start_llm_time = time.time()

                    # Extract keywords using LLM
                    keywords, categories = extract_keywords_with_llm(markdown_text, TAX_KEYWORDS, entry['section_number'], docx_filename)

                    # Calculate and print LLM processing time
                    llm_processing_time = time.time() - start_llm_time
                    print(f"LLM processing time for keywords extraction: {llm_processing_time:.2f} seconds")

                    # Adjust section to legislation + section numbering only
                    section_number_only = entry['section_number']

                    # Prepend full reference to text
                    markdown_text_with_ref = f"{full_reference}\n\n{markdown_text}"

                    # Add section number to keywords list
                    keywords_with_section = list(keywords)
                    if section_number_only not in keywords_with_section:
                        keywords_with_section.append(section_number_only)
                        logging.info(f"Added section number '{section_number_only}' to keywords")

                    chunk_metadata = {
                        "keywords": keywords_with_section,
                        "categories": list(categories),
                        "full_reference": full_reference,
                        "source_file": docx_filename,
                        "creation_date": datetime.now().isoformat()
                    }

                    chunks.append({
                        "section": section_number_only,
                        "text": markdown_text_with_ref,
                        "metadata": chunk_metadata
                    })
                
                chunks_count = len(chunks)
                
                # Save each chunk as individual JSON files
                docx_name = os.path.splitext(docx_filename)[0].replace(" ", "_").replace("-", "_")
                dt_string = datetime.now().strftime('%Y%m%d_%H%M%S')
                json_files_counter = 0
                
                # If in test mode, only process the first 20 chunks
                chunks_to_process = chunks[:20] if test_mode else chunks
                
                for chunk in chunks_to_process:
                    # Use the section field directly from the chunk
                    section_number = chunk["section"]
                    
                    # Make the section number safe for filenames
                    safe_section = section_number.replace(".", "_").replace("-", "_")
                    
                    # Generate unique filename starting with section number
                    json_filename = f"{safe_section}_{docx_name}_{dt_string}.json"
                    
                    # Check if this filename already exists
                    if os.path.exists(os.path.join(json_dir, json_filename)):
                        # Add a suffix to make it unique
                        json_filename = f"{safe_section}_{json_files_counter}_{docx_name}_{dt_string}.json"
                    
                    json_output_file = os.path.join(json_dir, json_filename)
                    try:
                        with open(json_output_file, "w", encoding="utf-8") as f:
                            json.dump(chunk, f, indent=2)
                        logging.info(f"Successfully saved JSON file: {json_output_file}")
                    except Exception as e:
                        logging.error(f"Failed to save JSON file {json_output_file}: {str(e)}")
                    json_files_counter += 1
                    print(f"Chunk for section {section_number} saved to {json_filename}")
                    
                    # Add debug output to show keywords for each JSON file in test mode
                    if test_mode:
                        print(f"Keywords for this chunk: {chunk['metadata']['keywords']}")
                
                if test_mode:
                    print(f"\n*** TEST MODE: Processed {json_files_counter} chunks out of {len(chunks)} total chunks ***")
                else:
                    print(f"Wrote {json_files_counter} actual JSON files out of {len(chunks)} chunks")
                
                # Write a CSV with all chunks
                csv_filename = f"{docx_name}_{dt_string}.csv"
                csv_output_file = os.path.join(json_dir, csv_filename)
                write_chunks_to_csv(chunks, csv_output_file)
                print(f"All chunks saved in CSV format to {csv_output_file}")
                
                file_end = time.time()
                elapsed = file_end - file_start

                # Convert elapsed time to hours, minutes, and seconds
                hours, rem = divmod(elapsed, 3600)
                minutes, seconds = divmod(rem, 60)

                summary_filename = f"{docx_name}_summary_{dt_string}.txt"
                summary_log_path = os.path.join(json_dir, summary_filename)
                write_individual_summary_log(summary_log_path, docx_filename, len(toc_entries), match_counter, chunks_count, elapsed, chunks)

                # Update the console output with both file time and overall time
                print(f"Completed processing {docx_filename} in {int(hours)}h {int(minutes)}m {int(seconds)}s")
                
                # Show elapsed time since start of run
                total_elapsed = time.time() - overall_start_time
                t_hours, t_rem = divmod(total_elapsed, 3600)
                t_minutes, t_seconds = divmod(t_rem, 60)
                print(f"Total time elapsed since start: {int(t_hours)}h {int(t_minutes)}m {int(t_seconds)}s")
                
                # Update the console output to show the new log location
                print(f"Summary log written to {summary_log_path}")
                
            except Exception as e:
                print(f"ERROR processing {docx_filename}: {str(e)}")
                import traceback
                traceback.print_exc()
                continue
        
    except Exception as e:
        print(f"ERROR: An exception occurred in main process: {str(e)}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()
